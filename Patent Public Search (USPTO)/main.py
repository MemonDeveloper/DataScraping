import requests
import time
import re
import os
import google.generativeai as genai
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from datetime import datetime
from docx2pdf import convert
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# Replace with your actual API key
GEMINI_API_KEY = "AIzaSyDkcuEXZczdLZf-Njvxa_7IJ4GBdXb_E5Y"  
# Configure API
genai.configure(api_key=GEMINI_API_KEY)
# Setup model
model = genai.GenerativeModel("models/gemini-2.0-flash")

# Define folder name
output_folder = "OCX_InfoSyncEngine"
# Create folder if it doesn't exist
os.makedirs(output_folder, exist_ok=True)

cookies = {
    '_ga': 'GA1.1.1490333756.1752502650',
    'QSI_SI_7WgrKZZwMjtuFh4_intercept': 'true',
    '_ga_Q5HVZK168H': 'GS2.1.s1752514811$o2$g1$t1752514900$j57$l0$h0',
    'aws-waf-token': '7410d63c-cc22-4642-8563-e0d0e2738959:BQoAqEqAINsuAAAA:7Z5tX16/G0wRLhrROmSckRcxC7xyYlasiJmgOc8vzsRqlkogqBY9Dcd/2m+45PHfz9NnhF41Lir5I4nj2s7ncPvwoxl5oISwygfFiuo7r1M3ro6GBbF7gDfWFgMYkeUv5u52IaPjpIdKufwNtFIZNUNsdMQWPRNPVdKyDEl53nnUvc5ARKm7ecyt+wuvskZuIi+iAGFp55UzcCXOq54m+OAysNyRjRhz1vJscSuWp/00mpZ4GHRA2pjEBhU12jQo',
    '_ga_M4L1KRPWXE': 'GS2.1.s1752517472$o1$g1$t1752517489$j43$l0$h0',
    '_ga_1KBHG37G8W': 'GS2.1.s1752514821$o3$g1$t1752517490$j60$l0$h0',
    '_ga_CD30TTEK1F': 'GS2.1.s1752600612$o6$g0$t1752600612$j60$l0$h0',
    '_ga_F4Q7EX1K95': 'GS2.1.s1752759228$o9$g1$t1752759254$j34$l0$h0',
    '_ga_CSLL4ZEK4L': 'GS2.1.s1752759229$o9$g1$t1752759255$j34$l0$h0',
}

headers = {
    'accept': '*/*',
    'accept-language': 'en-US,en;q=0.9',
    'cache-control': 'no-cache',
    'content-type': 'application/json; charset=UTF-8',
    'origin': 'https://ppubs.uspto.gov',
    'pragma': 'no-cache',
    'priority': 'u=1, i',
    'referer': 'https://ppubs.uspto.gov/pubwebapp/',
    'sec-ch-ua': '"Not)A;Brand";v="8", "Chromium";v="138", "Google Chrome";v="138"',
    'sec-ch-ua-mobile': '?0',
    'sec-ch-ua-platform': '"Windows"',
    'sec-fetch-dest': 'empty',
    'sec-fetch-mode': 'cors',
    'sec-fetch-site': 'same-origin',
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36',
    'x-access-token': 'eyJzdWIiOiI2OTQ5MjExNy1kMDRmLTRjMjEtYTdkYS1kMTU5M2JhZmM1MzUiLCJ2ZXIiOiI1ZDUzNjk1Mi1lYTU5LTRlZTMtYjVkOC0zNmQyNjQ0MGRkMzAiLCJleHAiOjB9',
    'x-requested-with': 'XMLHttpRequest',
    # 'cookie': '_ga=GA1.1.1490333756.1752502650; QSI_SI_7WgrKZZwMjtuFh4_intercept=true; _ga_Q5HVZK168H=GS2.1.s1752514811$o2$g1$t1752514900$j57$l0$h0; aws-waf-token=7410d63c-cc22-4642-8563-e0d0e2738959:BQoAqEqAINsuAAAA:7Z5tX16/G0wRLhrROmSckRcxC7xyYlasiJmgOc8vzsRqlkogqBY9Dcd/2m+45PHfz9NnhF41Lir5I4nj2s7ncPvwoxl5oISwygfFiuo7r1M3ro6GBbF7gDfWFgMYkeUv5u52IaPjpIdKufwNtFIZNUNsdMQWPRNPVdKyDEl53nnUvc5ARKm7ecyt+wuvskZuIi+iAGFp55UzcCXOq54m+OAysNyRjRhz1vJscSuWp/00mpZ4GHRA2pjEBhU12jQo; _ga_M4L1KRPWXE=GS2.1.s1752517472$o1$g1$t1752517489$j43$l0$h0; _ga_1KBHG37G8W=GS2.1.s1752514821$o3$g1$t1752517490$j60$l0$h0; _ga_CD30TTEK1F=GS2.1.s1752600612$o6$g0$t1752600612$j60$l0$h0; _ga_F4Q7EX1K95=GS2.1.s1752759228$o9$g1$t1752759254$j34$l0$h0; _ga_CSLL4ZEK4L=GS2.1.s1752759229$o9$g1$t1752759255$j34$l0$h0',
}

json_data1 = {
    'start': 0,
    'pageCount': 500,
    'sort': 'date_publ desc',
    'docFamilyFiltering': 'familyIdFiltering',
    'searchType': 1,
    'familyIdEnglishOnly': True,
    'familyIdFirstPreferred': 'US-PGPUB',
    'familyIdSecondPreferred': 'USPAT',
    'familyIdThirdPreferred': 'FPRS',
    'showDocPerFamilyPref': 'showEnglish',
    'queryId': 0,
    'tagDocSearch': False,
    'query': {
        'anchorDocIds': None,
        'querySource': 'brs',
        'caseId': 94259020,
        'hl_snippets': '2',
        'op': 'OR',
        'q': 'TTL ("system") AND ABST("data optimization" OR "information filtering")',
        'queryName': 'TTL ("system") AND ABST("data optimization" OR "information filtering")',
        'highlights': '1',
        'qt': 'brs',
        'spellCheck': False,
        'viewName': 'tile',
        'plurals': True,
        'britishEquivalents': True,
        'databaseFilters': [
            {
                'databaseName': 'US-PGPUB',
                'countryCodes': [],
            },
            {
                'databaseName': 'USPAT',
                'countryCodes': [],
            },
            {
                'databaseName': 'USOCR',
                'countryCodes': [],
            },
        ],
        'searchType': 1,
        'ignorePersist': True,
        'userEnteredQuery': 'TTL ("system") AND ABST("data optimization" OR "information filtering")',
    },
}

json_data2 = {
    'start': 0,
    'pageCount': 500,
    'sort': 'date_publ desc',
    'docFamilyFiltering': 'familyIdFiltering',
    'searchType': 1,
    'familyIdEnglishOnly': True,
    'familyIdFirstPreferred': 'US-PGPUB',
    'familyIdSecondPreferred': 'USPAT',
    'familyIdThirdPreferred': 'FPRS',
    'showDocPerFamilyPref': 'showEnglish',
    'queryId': 0,
    'tagDocSearch': False,
    'query': {
        'anchorDocIds': None,
        'querySource': 'brs',
        'caseId': 94259020,
        'hl_snippets': '2',
        'op': 'OR',
        'q': 'TTL ("platform") AND ABST("interactive engine" OR "multi-layer")',
        'queryName': 'TTL ("platform") AND ABST("interactive engine" OR "multi-layer")',
        'highlights': '1',
        'qt': 'brs',
        'spellCheck': False,
        'viewName': 'tile',
        'plurals': True,
        'britishEquivalents': True,
        'databaseFilters': [
            {
                'databaseName': 'US-PGPUB',
                'countryCodes': [],
            },
            {
                'databaseName': 'USPAT',
                'countryCodes': [],
            },
            {
                'databaseName': 'USOCR',
                'countryCodes': [],
            },
        ],
        'searchType': 1,
        'ignorePersist': True,
        'userEnteredQuery': 'TTL ("platform") AND ABST("interactive engine" OR "multi-layer")',
    },
}

json_data3 = {
    'start': 0,
    'pageCount': 500,
    'sort': 'date_publ desc',
    'docFamilyFiltering': 'familyIdFiltering',
    'searchType': 1,
    'familyIdEnglishOnly': True,
    'familyIdFirstPreferred': 'US-PGPUB',
    'familyIdSecondPreferred': 'USPAT',
    'familyIdThirdPreferred': 'FPRS',
    'showDocPerFamilyPref': 'showEnglish',
    'queryId': 0,
    'tagDocSearch': False,
    'query': {
        'anchorDocIds': None,
        'querySource': 'brs',
        'caseId': 94259020,
        'hl_snippets': '2',
        'op': 'OR',
        'q': 'TTL("module") AND ABST("coordination protocol") AND ABST("logic management")',
        'queryName': 'TTL("module") AND ABST("coordination protocol") AND ABST("logic management")',
        'highlights': '1',
        'qt': 'brs',
        'spellCheck': False,
        'viewName': 'tile',
        'plurals': True,
        'britishEquivalents': True,
        'databaseFilters': [
            {
                'databaseName': 'US-PGPUB',
                'countryCodes': [],
            },
            {
                'databaseName': 'USPAT',
                'countryCodes': [],
            },
            {
                'databaseName': 'USOCR',
                'countryCodes': [],
            },
        ],
        'searchType': 1,
        'ignorePersist': True,
        'userEnteredQuery': 'TTL("module") AND ABST("coordination protocol") AND ABST("logic management")',
    },
}

# JSON queries with their names
all_json_queries = [
    (json_data1, json_data1['query']['queryName']),
    (json_data2, json_data2['query']['queryName']),
    (json_data3, json_data3['query']['queryName'])
]

all_ids = set()          # ensures uniqueness
doc_type_pairs = []      # stores (documentId, type, query_name)

# STEP 1: Fetch documentId and type
for idx, (json_data, query_name) in enumerate(all_json_queries, start=1):
    print(f"🔍 Sending query {idx}...")
    response = requests.post("https://ppubs.uspto.gov/api/searches/searchWithBeFamily", json=json_data, headers=headers, cookies=cookies)

    if response.status_code == 200:
        data = response.json()
        patents = data.get('patents', [])
        
        doc_ids_this_round = []
        for doc in patents:
            if 'documentId' in doc and 'type' in doc:
                doc_id = doc['documentId'].replace(" ", "-")
                doc_type = doc['type']
                doc_type_pairs.append((doc_id, doc_type, query_name))
                doc_ids_this_round.append(doc_id)
        
        # Use top 100 document IDs from current result
        top_100_ids = doc_ids_this_round[:100]
        all_ids.update(top_100_ids)
    else:
        print(f"❌ Query {idx} failed: {response.status_code} - {response.text}")

# STEP 2: Convert to sorted list of unique IDs
unique_ids = sorted(all_ids)
# STEP 3: Sort doc_type_pairs based on query_name order
query_order = {query_name: i for i, (_, query_name) in enumerate(all_json_queries)}
doc_type_pairs = sorted(doc_type_pairs, key=lambda x: query_order.get(x[2], float('inf')))

# Output results
print(f"\n✅ Total unique document IDs: {len(unique_ids)}")

def get_gemini_summary(abstract: str, claims: str) -> str:
    combined_text = f"""You are a patent summarization expert.

Summarize the following patent based on its abstract and claims in **2-3 concise sentences**. Avoid legal jargon and keep it readable for non-technical audiences.

Abstract:
{abstract}

Claims:
{claims}
"""
    try:
        response = model.generate_content(combined_text)
        return response.text.strip()
    except Exception as e:
        print(f"❌ Gemini summary failed: {e}")
        return "Summary generation failed."

def clean_html_with_spans(text):
    """
    Cleans patent content while preserving:
    - Paragraph numbers like [0025]
    - Bold elements (<strong> or <b> tags)
    - Mathematical formulas
    - Proper line breaks
    """
    if not text or not isinstance(text, str):
        return text

    # Convert HTML tags to consistent format
    text = text.replace('<b>', '<strong>').replace('</b>', '</strong>')
    
    # Process with BeautifulSoup
    soup = BeautifulSoup(text, 'html.parser')
    
    # Handle paragraph numbers (like [0025])
    for para in soup.find_all(string=re.compile(r'\[\d+\]')):
        new_text = re.sub(r'\[(\d+)\]', r'[\1] ', para)
        para.replace_with(new_text)
    
    # Remove unwanted tags but keep content (except strong and math tags)
    for tag in soup.find_all(['figref', 'table-wrap', 'crossref', 'span']):
        if not tag.find_parents('math'):  # Don't unwrap spans inside math tags
            tag.unwrap()
        
    # Preserve strong tags for bold formatting (remove attributes only)
    for strong in soup.find_all('<strong>'):
        strong.replace_with('')
    
    for strong in soup.find_all('</strong>. '):
        strong.replace_with('')

    # Handle line breaks
    for br in soup.find_all('br'):
        br.replace_with('\n')

    # Special handling for math elements
    for math in soup.find_all('math'):
        # Create a clean representation of the math content
        math_id = math.get('id', '')
        math_num = math.get('num', '')
        
        # Extract the actual math content
        math_content = math.get_text(' ', strip=True)
        math_content = re.sub(r'\s+', ' ', math_content)  # Normalize whitespace
        
        # Create a readable label
        math_label = f"[Math Formula {math_num}]" if math_num else "[Math Formula]"
        
        # Replace with both label and content for better readability
        math.replace_with(f"{math_label}: {math_content}")

    # Final cleanup of whitespace
    cleaned_text = str(soup)
    cleaned_text = re.sub(r'\n\s+\n', '\n\n', cleaned_text)  # Reduce multiple newlines
    cleaned_text = re.sub(r'(\S)\n(\S)', r'\1 \2', cleaned_text)  # Fix broken words
    
    return cleaned_text

# --- Flatten JSON ignoring nulls ---
def flatten_json(y, parent_key='', sep='.'):
    items = []
    for k, v in y.items():
        if v is None:
            continue
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_json(v, new_key, sep=sep).items())
        elif isinstance(v, list):
            if all(isinstance(i, dict) for i in v):
                for idx, item in enumerate(v):
                    items.extend(flatten_json(item, f"{new_key}[{idx}]", sep=sep).items())
            else:
                cleaned_list = [str(i) for i in v if i is not None]
                if cleaned_list:
                    items.append((new_key, ', '.join(cleaned_list)))
        else:
            items.append((new_key, v))
    return dict(items)

# --- Fetch patent details ---
def fetch_patent_details(doc_id, doc_type):
    global success_count, fail_count
    detail_url = f"https://ppubs.uspto.gov/api/patents/highlight/{doc_id}"
    params = {
        'queryId': '87362852',
        'source': doc_type,
        'includeSections': 'true',
        'uniqueId': ''
    }
    try:
        response = requests.get(detail_url, params=params, cookies=cookies, headers=headers)
        if response.status_code == 200:
            return response.json()
        else:
            return None
    except Exception as e:
        print(f"❌ Exception occurred: {e}")
        return None

# --- Init PDF & DOCX ---
doc = Document()
doc.add_heading('Patent Details Report', 0)

success_count = 0
fail_count = 0
current_query = None
first_patent_in_query = True

for i, (doc_id, doc_type, query_name) in enumerate(doc_type_pairs[:5], start=1):
    # Check if we're starting a new query section
    if query_name != current_query:
        current_query = query_name
        first_patent_in_query = True
    
    result = fetch_patent_details(doc_id, doc_type)
    if not result:
        fail_count += 1
        print(f"❌ Failed #{fail_count}: {doc_id} ({doc_type})")
        continue

    # Clean known HTML fields before flattening
    for html_key in ["abstractHtml", "descriptionHtml", "claimsHtml", "backgroundTextHtml"]:
        if html_key in result:
            result[html_key] = clean_html_with_spans(result[html_key])

    flat_data = flatten_json(result)
    success_count += 1
    patent_number = flat_data.get("documentId", "")
    title = flat_data.get("inventionTitle", "Patent Details")

    # --- DOCX Output ---
    try:
        # Add a page break before each new patent
        if not first_patent_in_query:
            doc.add_page_break()

        # Add query header
        query_para = doc.add_paragraph(f"QUERY: {current_query.upper()}", style='Heading 1')
        for run in query_para.runs:
            run.font.name = 'Aptos Narrow'
            run.font.size = Pt(16)
            run.bold = True
                
        # Patent number with HTML span handling
        patent_number_clean = clean_html_with_spans(patent_number)
        title_para = doc.add_paragraph()
        title_label = title_para.add_run("PATENT NUMBER: ")
        title_label.font.name = 'Aptos Narrow'
        title_label.font.size = Pt(14)
        title_label.bold = True
        
        # Add patent number with preserved formatting but without span attributes
        soup = BeautifulSoup(patent_number_clean, 'html.parser')
        for content in soup.contents:
            if content.name == 'span':
                text_run = title_para.add_run(content.get_text())
                text_run.font.name = 'Aptos Narrow'
                text_run.font.size = Pt(12)
            else:
                text_run = title_para.add_run(str(content))
                text_run.font.name = 'Aptos Narrow'
                text_run.font.size = Pt(12)
        
        # Patent title with HTML span handling
        title_clean = clean_html_with_spans(title)
        title_para = doc.add_paragraph()
        title_label = title_para.add_run("TITLE: ")
        title_label.font.name = 'Aptos Narrow'
        title_label.font.size = Pt(14)
        title_label.bold = True
        
        soup = BeautifulSoup(title_clean, 'html.parser')
        for content in soup.contents:
            if content.name == 'span':
                text_run = title_para.add_run(content.get_text())
                text_run.font.name = 'Aptos Narrow'
                text_run.font.size = Pt(12)
            else:
                text_run = title_para.add_run(str(content))
                text_run.font.name = 'Aptos Narrow'
                text_run.font.size = Pt(12)
        
        # Patent details with HTML span handling
        for k, v in flat_data.items():
            if k in ["documentId", "inventionTitle"]:
                continue
                
            para = doc.add_paragraph()
            
            # Label in UPPERCASE, Aptos Narrow 14pt Bold
            label = para.add_run(f"{k.upper()}: ")
            label.font.name = 'Aptos Narrow'
            label.font.size = Pt(14)
            label.bold = True
            
            # Value with HTML span handling
            if isinstance(v, str):
                v_clean = clean_html_with_spans(v)
                soup = BeautifulSoup(v_clean, 'html.parser')
                for content in soup.contents:
                    if content.name == 'span':
                        text_run = para.add_run(content.get_text())
                        text_run.font.name = 'Aptos Narrow'
                        text_run.font.size = Pt(12)
                    else:
                        text_run = para.add_run(str(content))
                        text_run.font.name = 'Aptos Narrow'
                        text_run.font.size = Pt(12)
            else:
                value = para.add_run(str(v))
                value.font.name = 'Aptos Narrow'
                value.font.size = Pt(12)
        
        abstract_text = clean_html_with_spans(result.get("abstractHtml", ""))
        claims_text = clean_html_with_spans(result.get("claimsHtml", ""))
        summary = get_gemini_summary(abstract_text, claims_text)

        summary_para = doc.add_paragraph()
        summary_label = summary_para.add_run("GEMINI SUMMARY: ")
        summary_label.font.name = 'Aptos Narrow'
        summary_label.font.size = Pt(14)
        summary_label.bold = True

        summary_text = summary_para.add_run(summary)
        summary_text.font.name = 'Aptos Narrow'
        summary_text.font.size = Pt(12)

        # Add separator
        doc.add_paragraph("-" * 100)
        
    except Exception as e:
        print(f"❌ DOCX generation failed for {doc_id}: {e}")
    
    first_patent_in_query = False
    print(f"✅ Success #{success_count}: {doc_id} ({doc_type})")
    time.sleep(1)

# --- Save files ---
timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
output_folder = "OCX_InfoSyncEngine"
os.makedirs(output_folder, exist_ok=True)

docx_path = os.path.join(output_folder, f"OCX_InfoSyncEngine_{timestamp}.docx")
pdf_path = os.path.join(output_folder, f"OCX_InfoSyncEngine_{timestamp}.pdf")

# Save DOCX
doc.save(docx_path)
print(f"✅ DOCX file saved: {docx_path}")

# Convert to PDF
try:
    convert(docx_path, pdf_path)
    print(f"✅ PDF file saved: {pdf_path}")
except Exception as e:
    print(f"❌ Failed to convert to PDF: {e}")

print("\n📊 Fetch Summary:")
print(f"✅ Total Success: {success_count}")
print(f"❌ Total Failed : {fail_count}")

# -- Setup --
SERVICE_ACCOUNT_FILE = 'service_account.json'
SCOPES = ['https://www.googleapis.com/auth/drive.file']
FOLDER_NAME = "OCX_InfoSyncEngine"

credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES
)
drive_service = build('drive', 'v3', credentials=credentials)

# --- Step 1: Get or create folder ---
def get_or_create_folder(name):
    # Search for folder
    query = f"name = '{name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
    items = results.get('files', [])

    if items:
        folder_id = items[0]['id']
        print(f"📂 Folder found: {name} (ID: {folder_id})")
    else:
        # Create folder
        file_metadata = {
            'name': name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        file = drive_service.files().create(body=file_metadata, fields='id').execute()
        folder_id = file.get('id')
        print(f"🆕 Folder created: {name} (ID: {folder_id})")
    
    return folder_id

# --- Step 2: Upload file ---
def upload_file(file_path, parent_folder_id):
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [parent_folder_id]
    }
    media = MediaFileUpload(file_path, resumable=True)
    uploaded_file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id'
    ).execute()
    print(f"☁️ Uploaded: {file_path} → Drive ID: {uploaded_file['id']}")

# --- Call from your existing script ---
folder_id = get_or_create_folder(FOLDER_NAME)

# Replace these with your actual file paths
upload_file(docx_path, folder_id)
upload_file(pdf_path, folder_id)
