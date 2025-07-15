import requests
import time
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF
from datetime import datetime
import os

# Define folder name
output_folder = "OCX_InfoSyncEngine"

# Create folder if it doesn't exist
os.makedirs(output_folder, exist_ok=True)

cookies = {
    '_ga': 'GA1.1.1490333756.1752502650',
    'QSI_SI_7WgrKZZwMjtuFh4_intercept': 'true',
    '_ga_CD30TTEK1F': 'GS2.1.s1752514458$o5$g0$t1752514458$j60$l0$h0',
    '_ga_Q5HVZK168H': 'GS2.1.s1752514811$o2$g1$t1752514900$j57$l0$h0',
    'aws-waf-token': '7410d63c-cc22-4642-8563-e0d0e2738959:BQoAqEqAINsuAAAA:7Z5tX16/G0wRLhrROmSckRcxC7xyYlasiJmgOc8vzsRqlkogqBY9Dcd/2m+45PHfz9NnhF41Lir5I4nj2s7ncPvwoxl5oISwygfFiuo7r1M3ro6GBbF7gDfWFgMYkeUv5u52IaPjpIdKufwNtFIZNUNsdMQWPRNPVdKyDEl53nnUvc5ARKm7ecyt+wuvskZuIi+iAGFp55UzcCXOq54m+OAysNyRjRhz1vJscSuWp/00mpZ4GHRA2pjEBhU12jQo',
    '_ga_M4L1KRPWXE': 'GS2.1.s1752517472$o1$g1$t1752517489$j43$l0$h0',
    '_ga_1KBHG37G8W': 'GS2.1.s1752514821$o3$g1$t1752517490$j60$l0$h0',
    '_ga_F4Q7EX1K95': 'GS2.1.s1752595279$o6$g1$t1752595298$j41$l0$h0',
    '_ga_CSLL4ZEK4L': 'GS2.1.s1752595279$o6$g1$t1752595298$j41$l0$h0',
}

headers = {
    'accept': '*/*',
    'accept-language': 'en-US,en;q=0.9',
    'content-type': 'application/json; charset=UTF-8',
    'origin': 'https://ppubs.uspto.gov',
    'priority': 'u=1, i',
    'referer': 'https://ppubs.uspto.gov/pubwebapp/',
    'sec-ch-ua': '"Not)A;Brand";v="8", "Chromium";v="138", "Google Chrome";v="138"',
    'sec-ch-ua-mobile': '?0',
    'sec-ch-ua-platform': '"Windows"',
    'sec-fetch-dest': 'empty',
    'sec-fetch-mode': 'cors',
    'sec-fetch-site': 'same-origin',
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36',
    'x-access-token': 'eyJzdWIiOiJlODFlMzA1YS0xNjJiLTQ5NWUtOWRkZC05MjNiNzI0YmM1MmIiLCJ2ZXIiOiJhNWI2M2JjZS00MzFhLTRlODItYWFhOC00YmE5NzFmMjQwYmQiLCJleHAiOjB9',
    'x-requested-with': 'XMLHttpRequest',
    # 'cookie': '_ga=GA1.1.1490333756.1752502650; QSI_SI_7WgrKZZwMjtuFh4_intercept=true; _ga_CD30TTEK1F=GS2.1.s1752514458$o5$g0$t1752514458$j60$l0$h0; _ga_Q5HVZK168H=GS2.1.s1752514811$o2$g1$t1752514900$j57$l0$h0; aws-waf-token=7410d63c-cc22-4642-8563-e0d0e2738959:BQoAqEqAINsuAAAA:7Z5tX16/G0wRLhrROmSckRcxC7xyYlasiJmgOc8vzsRqlkogqBY9Dcd/2m+45PHfz9NnhF41Lir5I4nj2s7ncPvwoxl5oISwygfFiuo7r1M3ro6GBbF7gDfWFgMYkeUv5u52IaPjpIdKufwNtFIZNUNsdMQWPRNPVdKyDEl53nnUvc5ARKm7ecyt+wuvskZuIi+iAGFp55UzcCXOq54m+OAysNyRjRhz1vJscSuWp/00mpZ4GHRA2pjEBhU12jQo; _ga_M4L1KRPWXE=GS2.1.s1752517472$o1$g1$t1752517489$j43$l0$h0; _ga_1KBHG37G8W=GS2.1.s1752514821$o3$g1$t1752517490$j60$l0$h0; _ga_F4Q7EX1K95=GS2.1.s1752595279$o6$g1$t1752595298$j41$l0$h0; _ga_CSLL4ZEK4L=GS2.1.s1752595279$o6$g1$t1752595298$j41$l0$h0',
}

json_data1 = {
    'start': 0,
    'pageCount': 500,
    'sort': 'date_publ desc',
    'docFamilyFiltering': 'familyIdFiltering',
    'searchType': 1,
    'familyIdEnglishOnly': True,
    'familyIdFirstPreferred': 'US-PGPUB',
    'familyIdSecondPreferred': 'USPAT',
    'familyIdThirdPreferred': 'FPRS',
    'showDocPerFamilyPref': 'showEnglish',
    'queryId': 0,
    'tagDocSearch': False,
    'query': {
        'anchorDocIds': None,
        'querySource': 'brs',
        'caseId': 93845335,
        'hl_snippets': '2',
        'op': 'OR',
        'q': 'TTL ("system") AND ABST("data optimization" OR "information filtering")',
        'queryName': 'TTL ("system") AND ABST("data optimization" OR "information filtering")',
        'highlights': '1',
        'qt': 'brs',
        'spellCheck': False,
        'viewName': 'tile',
        'plurals': True,
        'britishEquivalents': True,
        'databaseFilters': [
            {
                'databaseName': 'US-PGPUB',
                'countryCodes': [],
            },
            {
                'databaseName': 'USPAT',
                'countryCodes': [],
            },
            {
                'databaseName': 'USOCR',
                'countryCodes': [],
            },
        ],
        'searchType': 1,
        'ignorePersist': True,
        'userEnteredQuery': 'TTL ("system") AND ABST("data optimization" OR "information filtering")',
    },
}

json_data2 = {
    'start': 0,
    'pageCount': 500,
    'sort': 'date_publ desc',
    'docFamilyFiltering': 'familyIdFiltering',
    'searchType': 1,
    'familyIdEnglishOnly': True,
    'familyIdFirstPreferred': 'US-PGPUB',
    'familyIdSecondPreferred': 'USPAT',
    'familyIdThirdPreferred': 'FPRS',
    'showDocPerFamilyPref': 'showEnglish',
    'queryId': 0,
    'tagDocSearch': False,
    'query': {
        'anchorDocIds': None,
        'querySource': 'brs',
        'caseId': 93845335,
        'hl_snippets': '2',
        'op': 'OR',
        'q': 'TTL ("platform") AND ABST("interactive engine" OR "multi-layer")',
        'queryName': 'TTL ("platform") AND ABST("interactive engine" OR "multi-layer")',
        'highlights': '1',
        'qt': 'brs',
        'spellCheck': False,
        'viewName': 'tile',
        'plurals': True,
        'britishEquivalents': True,
        'databaseFilters': [
            {
                'databaseName': 'US-PGPUB',
                'countryCodes': [],
            },
            {
                'databaseName': 'USPAT',
                'countryCodes': [],
            },
            {
                'databaseName': 'USOCR',
                'countryCodes': [],
            },
        ],
        'searchType': 1,
        'ignorePersist': True,
        'userEnteredQuery': 'TTL ("platform") AND ABST("interactive engine" OR "multi-layer")',
    },
}

json_data3 = {
    'start': 0,
    'pageCount': 500,
    'sort': 'date_publ desc',
    'docFamilyFiltering': 'familyIdFiltering',
    'searchType': 1,
    'familyIdEnglishOnly': True,
    'familyIdFirstPreferred': 'US-PGPUB',
    'familyIdSecondPreferred': 'USPAT',
    'familyIdThirdPreferred': 'FPRS',
    'showDocPerFamilyPref': 'showEnglish',
    'queryId': 0,
    'tagDocSearch': False,
    'query': {
        'anchorDocIds': None,
        'querySource': 'brs',
        'caseId': 93845335,
        'hl_snippets': '2',
        'op': 'OR',
        'q': 'TTL("module") AND ABST("coordination protocol") AND ABST("logic management")',
        'queryName': 'TTL("module") AND ABST("coordination protocol") AND ABST("logic management")',
        'highlights': '1',
        'qt': 'brs',
        'spellCheck': False,
        'viewName': 'tile',
        'plurals': True,
        'britishEquivalents': True,
        'databaseFilters': [
            {
                'databaseName': 'US-PGPUB',
                'countryCodes': [],
            },
            {
                'databaseName': 'USPAT',
                'countryCodes': [],
            },
            {
                'databaseName': 'USOCR',
                'countryCodes': [],
            },
        ],
        'searchType': 1,
        'ignorePersist': True,
        'userEnteredQuery': 'TTL("module") AND ABST("coordination protocol") AND ABST("logic management")',
    },
}

# JSON queries (already defined)
all_json_queries = [json_data1, json_data2, json_data3]

all_ids = set()          # ensures uniqueness
doc_type_pairs = []      # stores (documentId, type)

# STEP 1: Fetch documentId and type
for idx, json_data in enumerate(all_json_queries, start=1):
    print(f"🔍 Sending query {idx}...")
    response = requests.post("https://ppubs.uspto.gov/api/searches/searchWithBeFamily", json=json_data, headers=headers, cookies=cookies)

    if response.status_code == 200:
        data = response.json()
        patents = data.get('patents', [])
        
        doc_ids_this_round = []
        for doc in patents:
            if 'documentId' in doc and 'type' in doc:
                doc_id = doc['documentId'].replace(" ", "-")
                doc_type = doc['type']
                doc_type_pairs.append((doc_id, doc_type))
                doc_ids_this_round.append(doc_id)
        
        # Use top 100 document IDs from current result
        top_100_ids = doc_ids_this_round[:100]
        all_ids.update(top_100_ids)
    else:
        print(f"❌ Query {idx} failed: {response.status_code} - {response.text}")

# Convert to sorted list
unique_ids = sorted(all_ids)

# Output result
print(f"\n✅ Total unique document IDs: {len(unique_ids)}")

# --- STEP 2: Define function to fetch patent details ---
def fetch_patent_details(doc_id, doc_type):
    global success_count, fail_count
    detail_url = f"https://ppubs.uspto.gov/api/patents/highlight/{doc_id}"
    params = {
        'queryId': '87362852',
        'source': doc_type,
        'includeSections': 'true',
        'uniqueId': ''
    }

    try:
        response = requests.get(detail_url, params=params, cookies=cookies, headers=headers)
        if response.status_code == 200:
            return response.json()
        else:
            return None
    except Exception as e:
        print(f"❌ Exception occurred: {e}")
        return None

# Step 3: Create PDF and DOCX reports
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
pdf_filename = f"patent_details_{timestamp}.pdf"
docx_filename = f"patent_details_{timestamp}.docx"

pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()

doc = Document()
doc.add_heading('Patent Details Report', 0)

# Function to generate PDF and add to DOCX
def add_patent_to_documents(patent_data):
    patent_number = patent_data.get("documentId", "")
    app_number = patent_data.get("applicationNumber", "N/A")
    title = patent_data.get("inventionTitle", "N/A")
    abstract = patent_data.get("abstractHtml", "").replace("<br />", "").strip()
    filing_date = patent_data.get("applicationFilingDate", ["N/A"])[0][:10]
    inventor_names = "; ".join(patent_data.get("inventorsName", [])) or "N/A"

    raw_claims = patent_data.get("claimsHtml", "")
    if raw_claims:
        soup = BeautifulSoup(raw_claims, "html.parser")
        text = soup.get_text(separator="\n")
        claims_text = re.sub(r'\n+', '\n', text).strip()
    else:
        claims_text = "N/A"

    # --- Create individual PDF ---
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Add and use Unicode font
        pdf.add_font('ArialUnicode', '', 'ArialUnicodeMS.ttf', uni=True)
        pdf.set_font('ArialUnicode', '', 16)

        # Header (Patent Number)
        pdf.set_text_color(0)
        pdf.cell(0, 10, f"Patent Number: {patent_number}", ln=True)

        # Body font (normal size)
        pdf.set_font('ArialUnicode', '', 10)

        # Main content
        pdf.multi_cell(0, 6, f"""Application No.: {app_number}
Filing Date: {filing_date}
Inventors: {inventor_names}

Abstract: {abstract}

Claims:
{claims_text}
""")

        # Summary placeholder
        pdf.multi_cell(0, 6, "\nSummary:\nGemini Summary Placeholder")

        # Save with safe filename
        safe_filename = patent_number.replace("/", "-").replace(":", "-").replace(" ", "-")
        pdf.output(os.path.join(output_folder, f"OCX_InfoSyncEngine_{safe_filename}.pdf"))
    except Exception as e:
        print(f"❌ PDF generation failed for {patent_number}: {e}")

    # --- Create individual DOCX ---
    try:
        doc = Document()
        doc.add_heading(title, level=1)

        info_items = [
            ("Patent Number:", patent_number),
            ("Application No.:", app_number),
            ("Filing Date:", filing_date),
            ("Inventors:", inventor_names),
            ("Abstract:", abstract),
            ("Claims:", claims_text)
        ]

        for label, value in info_items:
            paragraph = doc.add_paragraph()
            
            run_label = paragraph.add_run(f"{label} ")
            run_label.bold = True
            run_label.font.size = Pt(16)

            run_value = paragraph.add_run(f"{value}")
            run_value.bold = False
            run_value.font.size = Pt(16)

            # Add spacing after Abstract and Claims
            if label in ["Abstract:", "Claims:"]:
                doc.add_paragraph("")  # Add blank line

        # Add Summary section
        summary_para = doc.add_paragraph()
        summary_title = summary_para.add_run("Summary:\n")
        summary_title.bold = True
        summary_title.font.size = Pt(16)

        summary_body = summary_para.add_run("Gemini Summary Placeholder")
        summary_body.font.size = Pt(16)

        # Add horizontal separator
        doc.add_paragraph("-" * 100)

        # Save file
        safe_filename = patent_number.replace("/", "-").replace(":", "-").replace(" ", "-")
        doc.save(os.path.join(output_folder, f"OCX_InfoSyncEngine_{safe_filename}.docx"))

    except Exception as e:
        print(f"❌ DOCX generation failed for {patent_number}: {e}")

# Process all patents
success_count = 0
fail_count = 0

for i, (doc_id, doc_type) in enumerate(doc_type_pairs, start=1):
    result = fetch_patent_details(doc_id, doc_type)
    if result:
        success_count += 1
        add_patent_to_documents(result)
        print(f"✅ Success #{success_count}: {doc_id} ({doc_type})")
    else:
        fail_count += 1
        print(f"❌ Failed #{fail_count}: {doc_id} ({doc_type})")
    time.sleep(1)

# Summary
print("\n📊 Fetch Summary:")
print(f"✅ Total Success: {success_count}")
print(f"❌ Total Failed : {fail_count}")
print(f"📄 Combined DOCX saved as: {docx_filename}")
